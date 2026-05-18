import streamlit as st
import pandas as pd
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
import urllib.parse

# --- KONFIGURASI GOOGLE SHEETS ---
URL_SHEETS = "https://script.google.com/macros/s/AKfycbzpj7xT43iJOxL-f9CksWBuUYswKHirEZ_i7mX3t0AWEdwy2R6OzOVRTsdftMnaTPPt/exec"

def save_to_sheets(data_laporan):
    try:
        response = requests.post(URL_SHEETS, json=data_laporan)
        return response.status_code == 200
    except: 
        return False

# --- 1. DATABASE CHECKLIST MAINTENANCE & SERVICE ---
if 'db_checklist' not in st.session_state:
    st.session_state.db_checklist = {
        "AeHealth": {
            "AERC-3": {
                "Bulanan": {
                    "Mati": ["Periksa kondisi Reagent Diluent & Lyse", "Rendam chamber RBC & WBC", "Cleaning Flow Cell", "Cleaning Selang Sampling / Tubing", "Cleaning Jarum Aspirate", "Cleaning Fan Pendingin", "Bersihkan Debu Pada Komponen Internal", "Cleaning Body Alat"],
                    "Hidup": ["Prime Diluent & Lyse", "Prime Seluruh Reagent", "Jalankan Self Check", "Test Background 2–3 kali", "Test Printer", "QC (Quality Control)"]
                },
                "Service": ["Periksa kondisi reagen Diluent", "Periksa kondisi Reagent Lyse", "Rendam chamber RBC", "Rendam chamber WBC", "Cleaning Flow Cell", "Cleaning Selang Sampling / Tubing", "Cleaning Jarum Aspirate", "Cleaning Fan Pendingin", "Bersihkan Debu Pada Komponen Internal", "Cleaning Body Alat", "Prime Diluent", "Prime Lyse", "Prime Seluruh Reagen", "Jalankan Self Check", "Test Background 2–3 kali", "Test Printer", "QC (Quality Control)"]
            }
        },
        "Zybio": {
            "Z5": {
                "Bulanan": {
                    "Mati": ["Periksa kondisi Reagent Diluent & Lyse", "Rendam chamber RBC & WBC", "Cleaning Flow Cell", "Cleaning Selang Sampling / Tubing", "Cleaning Jarum Aspirate", "Cleaning Fan Pendingin", "Bersihkan Debu Pada Komponen Internal", "Cleaning Body Alat"],
                    "Hidup": ["Prime Diluent & Lyse", "Prime Seluruh Reagent", "Jalankan Self Check", "Test Background 3 kali", "Test Printer", "QC (Quality Control)"]
                }
            },
            "Z3 / Z31 / Z3 CRP": {
                "Bulanan": {
                    "Mati": ["Periksa kondisi Reagent Diluent & Lyse", "Rendam chamber RBC & WBC", "Cleaning Flow Cell", "Cleaning Selang Sampling / Tubing", "Cleaning Jarum Aspirate", "Cleaning Fan Pendingin", "Bersihkan Debu Pada Komponen Internal", "Cleaning Body Alat"],
                    "Hidup": ["Prime Diluent & Lyse", "Prime Seluruh Reagent", "Jalankan Self Check", "Test Background 2–3 kali", "Test Printer", "QC (Quality Control)"]
                }
            }
        },
        "Porlak": {
            "PJH-B200": {
                "Bulanan": {
                    "Mati": ["Periksa Kondisi Air Secara Visual", "Periksa Kualitas Air Menggunakan TDS meter", "Cleaning Filter Air", "Cleaning Cuvet", "Cleaning Reagent tray", "Cleaning Jarum Sample", "Cleaning Jarum Washer", "Cleaning Mixer / Pengaduk", "Cleaning Body Alat"],
                    "Hidup": ["Reset System", "Air Purge 3 kali", "Detergent Wash", "Wash Cell All", "Cek Cell Blank (Range 20.000–60.000)", "QC (Quality Control)"]
                }
            },
            "PJH-B60 / PJHB-101": {
                "Bulanan": {
                    "Mati": ["Cleaning Cuvet", "Cleaning Selang Aspirasi", "Cleaning Inkubator", "Cleaning Fan Pendingin", "Bersihkan Debu Pada Komponen Internal", "Cleaning Body Alat"],
                    "Hidup": ["Washing / Cuci 10 Detik", "Cek Background", "Periksa Suhu Inkubator (37°C)", "Periksa Kondisi Lampu (Range 20.000–60.000)", "Test Printer", "QC (Quality Control)"]
                }
            },
            "PJH-U300": {
                "Bulanan": {
                    "Hidup": ["Periksa Liquid Box", "Periksa Waste Box", "Cleaning Belt Conveyor", "Cleaning Body Alat", "Bersihkan Debu pada Komponen Internal", "QC (Quality Control)"]
                }
            },
            "PJH-H360": {
                "Bulanan": {
                    "Mati": ["Periksa kondisi Reagent Diluent & Lyse", "Rendam Chamber RBC & WBC", "Cleaning Flow Cell", "Cleaning Selang Sampling / Tubing", "Cleaning Jarum Aspirate", "Cleaning Fan Pendingin", "Bersihkan Debu Pada Komponen Internal", "Cleaning Body Alat"],
                    "Hidup": ["Prime Diluent", "Prime Lyse", "Prime Seluruh Reagent", "Jalankan Self Check", "Test Background 3 Kali", "Test Printer", "QC (Quality Control)"]
                }
            },
            "PJH-H610": {
                "Bulanan": {
                    "Mati": ["Periksa kondisi Reagent Diluent & Lyse", "Rendam chamber RBC & WBC", "Cleaning Flow Cell", "Cleaning Selang Sampling / Tubing", "Cleaning Jarum Aspirate", "Cleaning Fan Pendingin", "Bersihkan Debu Pada Komponen Internal", "Cleaning Body Alat"],
                    "Hidup": ["Prime Diluent", "Prime Lyse", "Prime Seluruh Reagent", "Jalankan Self Check", "Test Background 3 Kali", "Test Printer", "QC (Quality Control)"]
                }
            }
        },
        "Caretium": {
            "XI-931": {
                "Bulanan": {
                    "Mati": ["Cleaning Jarum Sample", "Cleaning Selang", "Cleaning Fan Pendingin", "Cleaning Body Alat", "Check Kondisi Peristaltic Tube"],
                    "Hidup": ["Prime Reagen", "Test Calibration Slope", "Test Printer", "QC (Quality Control)"]
                }
            }
        },
        "Lamuno": {
            "Lamuno X / Lamuno Pro": {
                "Bulanan": {
                    "Hidup": ["Periksa QC Internal, Pastikan Tidak Expired", "Periksa kondisi Reagent", "Pastikan Nomor Lot Reagent Sesuai Dengan Data Pada Alat", "Test Printer", "Cleaning Body Alat", "QC (Quality Control)"]
                }
            }
        },
        "Biontech": {
            "ES-380": {
                "Bulanan": {
                    "Mati": ["Periksa Kondisi Air Secara Visual", "Perikas Kualitas Air Menggunakan TDS Meter", "Cleaning Filter Air", "Cleaning Cuvet", "Cleaning Reagent Tray", "Cleaning Jarum Sample", "Cleaning Jarum Washer", "Cleaning Mixer / Pengaduk", "Cleaning Body Alat"],
                    "Hidup": ["Reset System", "Air Purge 3 kali", "Detergent Wash", "Wash Cell All", "Cek Cell Blank (Range 20.000–60.000)", "QC (Quality Control)"]
                }
            },
            "ES-101": {
                "Bulanan": {
                    "Mati": ["Cleaning Cuvet", "Cleaning Selang Aspirasi", "Cleaning Inkubator", "Cleaning Fan Pendingin", "Bersihkan Debu Pada Komponen Internal", "Cleaning Body Alat"],
                    "Hidup": ["Washing / Cuci 10 Detik", "Cek Background", "Periksa Suhu Inkubator (37°C)", "Periksa Kondisi Lampu (Range 20.000–60.000)", "Test Printer", "QC (Quality Control)"]
                }
            }
        },
        "Singseng": {
            "RD6S": {
                "Bulanan": {
                    "Mati": ["Periksa kondisi reagen Diluent & Lyse", "Rendam chamber RBC & WBC", "Cleaning Flow Cell", "Cleaning Selang Sampling / Tubing", "Cleaning Jarum Aspirate", "Cleaning Fan Pendingin", "Bersihkan Debu Pada Komponen Internal", "Cleaning Body Alat"],
                    "Hidup": ["Prime Diluent", "Prime Lyse", "Prime Seluruh Reagent", "Jalankan Self Check", "Test Background 3 Kali", "Test Printer", "QC (Quality Control)"]
                }
            },
            "RD3S": {
                "Bulanan": {
                    "Mati": ["Periksa kondisi reagen Diluent & Lyse", "Rendam chamber RBC & WBC", "Cleaning Flow Cell", "Cleaning Selang Sampling / Tubing", "Cleaning Jarum Aspirate", "Cleaning Fan Pendingin", "Bersihkan Debu Pada Komponen Internal", "Cleaning Body Alat"],
                    "Hidup": ["Prime Diluent", "Prime Lyse", "Prime Seluruh Reagent", "Jalankan Self Check", "Test Background 3 Kali", "Test Printer", "QC (Quality Control)"]
                }
            }
        },
        "Bioway": {
            "BW-300": {
                "Bulanan": {
                    "Hidup": ["Periksa Liquid Box", "Periksa Waste Box", "Cleaning Belt Conveyor", "Cleaning Body Alat", "Bersihkan Debu pada Komponen Internal", "QC (Quality Control)"]
                }
            }
        },
        "Heto": {
            "HT-H360": {
                "Bulanan": {
                    "Mati": ["Periksa kondisi Reagent Diluent & Lyse", "Rendam chamber RBC & WBC", "Cleaning Flow Cell", "Cleaning Selang Sampling / Tubing", "Cleaning Jarum Aspirate", "Cleaning Fan Pendingin", "Bersihkan Debu Pada Komponen Internal", "Cleaning Body Alat"],
                    "Hidup": ["Prime Diluent", "Prime Lyse", "Prime Seluruh Reagent", "Jalankan Self Check", "Test Background 3 Kali", "Test Printer", "QC (Quality Control)"]
                }
            },
            "HT-B200": {
                "Bulanan": {
                    "Mati": ["Periksa Kondisi Air Secara Visual", "Perikas Kualitas Air Menggunakan TDS Meter", "Cleaning Filter Air", "Cleaning Cuvet", "Cleaning Reagent Tray", "Cleaning Jarum Sample", "Cleaning Jarum Washer", "Cleaning Mixer / Pengaduk", "Cleaning Body Alat"],
                    "Hidup": ["Reset System", "Air Purge 3 kali", "Detergent Wash", "Wash Cell All", "Cek Cell Blank (Range 20.000–60.000)", "QC (Quality Control)"]
                }
            },
            "HT-U300": {
                "Bulanan": {
                    "Hidup": ["Periksa Liquid Box", "Periksa Waste Box", "Cleaning Belt Conveyor", "Cleaning Body Alat", "Bersihkan Debu pada Komponen Internal", "QC (Quality Control)"]
                }
            },
            "HT-B60 / HT-B101": {
                "Bulanan": {
                    "Mati": ["Cleaning Cuvet", "Cleaning Selang Aspirasi", "Cleaning Inkubator", "Cleaning Fan Pendingin", "Bersihkan Debu Pada Komponen Internal", "Cleaning Body Alat"],
                    "Hidup": ["Washing / Cuci 10 Detik", "Cek Background", "Periksa Suhu Inkubator (37°C)", "Periksa Kondisi Lampu (Range 20.000–60.000)", "Test Printer", "QC (Quality Control)"]
                }
            }
        }
    }

if 'db_alat' not in st.session_state:
    st.session_state.db_alat = {
        "AeHealth": ["AERC-3"],
        "Biontech": ["ES-380", "ES-101"],
        "Bioway":   ["BW-300"],
        "Caretium": ["XI-931"],
        "Heto": ["HT-H360", "HT-B200", "HT-U300", "HT-B60 / HT-B101"],
        "Lamuno": ["Lamuno X / Lamuno Pro"],
        "Porlak": ["PJH-B200", "PJH-H360", "PJH-H610", "PJH-B60 / PJHB-101", "PJH-U300"],
        "Singseng": ["RD6S", "RD3S"],
        "Zybio": ["Z5", "Z3 / Z31 / Z3-CRP"]
    }

if 'list_teknisi' not in st.session_state:
    st.session_state.list_teknisi = ["Muhammad Edy Surya", "Muhammad Abdul Majid", "Ari Nugroho", "Leonardus Naldo Poju"]

# Inisialisasi awal database link agar datanya tersimpan konstan di session state
if 'link_drive_data' not in st.session_state:
    st.session_state.link_drive_data = {
        "Link Folder Utama Report": "https://drive.google.com/drive/folders/example1",
        "Link Reagent & Manual Book": "https://drive.google.com/drive/folders/example2",
        "Link Dokumentasi Service": "https://drive.google.com/drive/folders/example3"
    }

# --- 2. KONFIGURASI HALAMAN & STYLE CSS ---
st.set_page_config(page_title="REI E-Report Teknisi", layout="wide", page_icon="⚙️")

st.markdown("""
<link href="https://fonts.googleapis.com/icon?family=Material+Icons" rel="stylesheet">
<style>
    @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@400;500;600;700;800&display=swap');
    
    [data-testid="stSidebar"] div[data-testid="stWidgetLabel"] {
        display: none !important;
    }
    [data-testid="stSidebar"] i[data-testid="stRadioCustomButton"] {
        display: none !important;
    }
    
    [data-testid="stSidebar"] div[role="radiogroup"] label {
        background-color: rgba(255, 255, 255, 0.05) !important;
        border: 1px solid rgba(255, 255, 255, 0.1) !important;
        padding: 12px 20px !important;
        border-radius: 8px !important;
        margin-bottom: 10px !important;
        display: block !important;
        width: 100% !important;
        transition: all 0.3s ease !important;
        cursor: pointer !important;
    }
    
    [data-testid="stSidebar"] div[role="radiogroup"] label:hover {
        background-color: rgba(255, 255, 255, 0.12) !important;
        border-color: rgba(245, 158, 11, 0.5) !important;
    }
    
    [data-testid="stSidebar"] div[role="radiogroup"] label[data-checked="true"] {
        background: linear-gradient(135deg, #f59e0b 0%, #d97706 100%) !important;
        border-color: #f59e0b !important;
        box-shadow: 0 4px 15px rgba(245, 158, 11, 0.4) !important;
    }
    
    [data-testid="stSidebar"] div[role="radiogroup"] label[data-checked="true"] p {
        color: #0f172a !important;
        font-weight: 700 !important;
    }
    
    html, body, [data-testid="stAppViewContainer"] { 
        font-family: 'Poppins', sans-serif; 
    }
    
    div[data-testid="stVerticalBlockBorderWrapper"] {
        background-color: var(--st-metric-background-color, rgba(255, 255, 255, 0.04)) !important;
        border-radius: 12px !important; 
        padding: 25px !important;
        border: 1px solid var(--st-input-border-color, rgba(255, 255, 255, 0.1)) !important;
    }
    
    div[data-testid="stVerticalBlockBorderWrapper"] p, 
    div[data-testid="stVerticalBlockBorderWrapper"] label,
    div[data-testid="stVerticalBlockBorderWrapper"] div {
        color: var(--text-color) !important;
    }
    
    .header-banner {
        background: linear-gradient(135deg, rgba(217, 119, 6, 0.85) 0%, rgba(245, 158, 11, 0.9) 100%), 
                    url('https://png.pngtree.com/thumb_back/fh260/background/20200807/pngtree-sunset-cityscape-night-scene-with-office-building-background-image_388725.jpg');
        background-size: cover;
        background-position: center;
        border-radius: 12px; 
        padding: 25px 35px; 
        margin-bottom: 25px; 
        color: white !important;
        display: flex; 
        justify-content: space-between; 
        align-items: center;
        box-shadow: 0 10px 25px rgba(0,0,0,0.2);
        flex-wrap: wrap;
        gap: 15px;
    }
    
    @media (max-width: 768px) {
        .header-banner {
            flex-direction: column !important;
            align-items: flex-start !important;
            padding: 20px !important;
            gap: 15px !important;
        }
        .banner-flex { width: 100% !important; }
        .rei-text { font-size: 45px !important; margin-right: 15px !important; }
        .title-block { padding-left: 15px !important; }
        .ereport-text { font-size: 18px !important; }
        .teknisi-text { font-size: 26px !important; }
        .header-banner > div:last-child {
            text-align: left !important;
            border-left: none !important;
            border-top: 1px solid rgba(255, 255, 255, 0.3) !important;
            padding-left: 0 !important;
            padding-top: 10px !important;
            width: 100% !important;
        }
    } 
    
    .banner-flex { display: flex; align-items: center; }
    .rei-text { font-size: 65px; font-weight: 800; color: #ffffff; letter-spacing: -3px; line-height: 1; margin-right: 20px; }
    .title-block { border-left: 3px solid rgba(255,255,255,0.4); padding-left: 20px; }
    .ereport-text { font-size: 24px; font-weight: 700; color: #ffffff; margin: 0; line-height: 1.1; }
    .teknisi-text { font-size: 34px; font-weight: 800; color: #0f172a; margin: 0; line-height: 1; letter-spacing: 1px; }
    
    .section-title {
        border-left: 4px solid #f59e0b; padding-left: 12px; font-size: 20px; font-weight: 600; color: var(--text-color); margin-top: 25px; margin-bottom: 15px;
    }
    
    div.stButton > button:first-child {
        background: #f59e0b !important; color: #0f172a !important;
        font-weight: 700 !important; height: 50px !important; width: 100% !important; border-radius: 8px !important; font-size: 16px !important; border: none !important;
        box-shadow: 0 4px 15px rgba(245, 158, 11, 0.3);
    }
    div.stButton > button:first-child:hover { background: #d97706 !important; color: white !important; }
    
    .whatsapp-btn {
        background-color: #25d366; color: white; padding: 12px; border-radius: 6px;
        text-decoration: none; font-weight: bold; display: block; text-align: center; margin-top: 10px;
    }
</style>
""", unsafe_allow_html=True)

# --- 3. FUNGSI GENERATE DOCX ---
def generate_docx(checklist, tech, id_tech, cust, sn, note, merek, tipe, periode, pt, branch, foto_list, ttd_t, ttd_c):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("MAINTENANCE REPORT")
    title_run.bold = True
    title_run.font.size = Pt(14)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    meta_table = doc.add_table(rows=8, cols=2)
    meta_table.style = 'Table Grid'
    meta_table.autofit = False
    
    for row in meta_table.rows:
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.2)

    identitas_data = [
        ("Perusahaan", pt),
        ("Cabang", branch),
        ("Teknisi", tech),
        ("Periode", periode),
        ("Customer", cust if cust else "-"),
        ("Alat", f"{merek} {tipe}"),
        ("S/N", sn if sn else "-"),
        ("Tanggal", datetime.now().strftime('%d/%m/%Y'))
    ]

    for idx, (label, value) in enumerate(identitas_data):
        cell_label = meta_table.rows[idx].cells[0]
        cell_val = meta_table.rows[idx].cells[1]
        p_lbl = cell_label.paragraphs[0]
        p_lbl.add_run(f" {label}").bold = True
        p_val = cell_val.paragraphs[0]
        p_val.add_run(f" {value}")

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    chk_title_p = doc.add_paragraph()
    chk_title_run = chk_title_p.add_run("CHECKLIST MAINTENANCE")
    chk_title_run.bold = True

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'No'
    hdr_cells[1].text = 'Pekerjaan'
    hdr_cells[2].text = 'Hasil'
    
    hdr_cells[0].width = Inches(0.5)
    hdr_cells[1].width = Inches(4.5)
    hdr_cells[2].width = Inches(1.4)
    
    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for i, item in enumerate(checklist):
        row_cells = table.add_row().cells
        langkah_bersih = item['Langkah']
        langkah_bersih = langkah_bersih.replace("(Mati)", "(Power OFF)").replace("(Hidup)", "(Power ON)")
        
        row_cells[0].text = str(i + 1)
        row_cells[1].text = f" {langkah_bersih}"
        row_cells[2].text = str(item['Hasil'])
        
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        row_cells[0].width = Inches(0.5)
        row_cells[1].width = Inches(4.5)
        row_cells[2].width = Inches(1.4)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    note_title_p = doc.add_paragraph()
    note_title_run = note_title_p.add_run("CATATAN")
    note_title_run.bold = True
    
    note_table = doc.add_table(rows=1, cols=1)
    note_table.style = 'Table Grid'
    note_table.rows[0].cells[0].width = Inches(6.4)
    
    p_note_box = note_table.rows[0].cells[0].paragraphs[0]
    p_note_box.add_run(f" {note}" if note else " ")
    note_table.rows[0].height = Inches(0.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    ttd_table = doc.add_table(rows=3, cols=2)
    ttd_table.autofit = False
    ttd_table.rows[0].cells[0].width = Inches(3.2)
    ttd_table.rows[0].cells[1].width = Inches(3.2)
    
    p_tek_t = ttd_table.rows[0].cells[0].paragraphs[0]
    p_tek_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tek_t.add_run("Technician,").font.size = Pt(10)
    
    p_cus_t = ttd_table.rows[0].cells[1].paragraphs[0]
    p_cus_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cus_t.add_run("Customer Approval,").font.size = Pt(10)
    
    ttd_table.rows[1].height = Inches(0.8)
    if ttd_t:
        p_img_t = ttd_table.rows[1].cells[0].paragraphs[0]
        p_img_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img_t.add_run().add_picture(ttd_t, width=Inches(1.3))
    if ttd_c:
        p_img_c = ttd_table.rows[1].cells[1].paragraphs[0]
        p_img_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img_c.add_run().add_picture(ttd_c, width=Inches(1.3))
        
    p_tek_n = ttd_table.rows[2].cells[0].paragraphs[0]
    p_tek_n.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tek_n.add_run(f"( {tech} )")
    
    p_cus_n = ttd_table.rows[2].cells[1].paragraphs[0]
    p_cus_n.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cus_n.add_run(f"( {cust if cust else '........................'} )")

    if foto_list:
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(20)
        
        h_foto = doc.add_heading('DOKUMENTASI FOTO', level=1)
        h_foto.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        table_foto = doc.add_table(rows=0, cols=2)
        row_cells = None
        for i, foto in enumerate(foto_list):
            if i % 2 == 0:
                row_cells = table_foto.add_row().cells
            cell = row_cells[i % 2]
            p_foto = cell.paragraphs[0]
            p_foto.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_foto = p_foto.add_run()
            run_foto.add_picture(foto, width=Inches(2.8))
            
    target = BytesIO()
    doc.save(target)
    return target.getvalue()

# --- 4. NAVIGATION & HEADER ---
with st.sidebar:
    st.image("https://rajaerba.id/wp-content/uploads/2022/09/logo-rei.png", width=160)
    st.markdown("""
    <div style="padding: 10px 0px 20px 0px; margin-top: -40px; border-bottom: 1px solid rgba(255,255,255,0.1); margin-bottom: 20px;">
        <h2 style="color: white; font-weight: 800; letter-spacing: -1px; margin: 0; font-size: 28px; line-height:1;">
            REI E-REPORT
        </h2>
        <p style="color: #f59e0b; font-size: 11px; margin: 5px 0 0 0; font-weight: 600; letter-spacing: 1.5px;">
            TECHNICAL SERVICE SYSTEM
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("### MAIN MENU")
    menu = st.radio("NAVIGASI", ["Maintenance", "Service", "Pengaturan"], label_visibility="collapsed")
    st.markdown("---")
    st.markdown("### DATA TEKNISI")
    nama_teknisi = st.selectbox("Nama Teknisi", st.session_state.list_teknisi)
    id_teknisi = st.text_input("Jabatan", "TEKNISI")
    nama_pt = st.text_input("Perusahaan", "PT. RAJAERBA INDOCHEM")
    lokasi_pt = st.selectbox("Cabang", ["Tangerang", "Semarang", "Jakarta", "Medan"])

st.markdown("""
<div class="header-banner">
    <div class="banner-flex">
        <div class="rei-text">REI</div>
        <div class="title-block">
            <p class="ereport-text">E-REPORT</p>
            <p class="teknisi-text">TEKNISI</p>
        </div>
    </div>
    <div style="text-align: right; font-size: 13px; border-left: 1px solid rgba(255,255,255,0.2); padding-left: 20px; color: #ffffff !important;">
        <span style="font-weight: 700; color: #0f172a;">PT. RAJAERBA INDOCHEM</span><br>
        <span style="opacity: 0.9; font-size: 11px;">Integritas • Akurat • Andal</span>
    </div>
</div>
""", unsafe_allow_html=True)

# --- 5. MAIN FORM LOGIC ---
if menu in ["Maintenance", "Service"]:
    st.markdown(f'<div class="section-title">Form {menu}</div>', unsafe_allow_html=True)
    
    if menu == "Maintenance":
        periode = st.selectbox("Periode Maintenance:", ["Bulanan", "3 Bulan", "6 Bulan", "12 Bulan"])
    else:
        periode = "Service"
    
    with st.container(border=True):
        st.markdown('<div class="section-title">Unit & Customer</div>', unsafe_allow_html=True)
        c1, c2 = st.columns(2, gap="small")
        with c1:
            merek = st.selectbox("Merek Alat", list(st.session_state.db_alat.keys()))
            tipe = st.selectbox("Tipe Alat", st.session_state.db_alat[merek])
            sn_alat = st.text_input("S/N Alat", placeholder="Masukkan Serial Number Alat")
        with c2:
            nama_customer = st.text_input("Nama Customer", placeholder="Masukkan Nama Customer")
            alamat_customer = st.text_input("Alamat", placeholder="Masukkan Alamat Customer")
            nomor_telepon = st.text_input("Nomor Telepon", placeholder="Masukkan Nomor Telepon")

    data_alat = st.session_state.db_checklist.get(merek, {}).get(tipe, {}).get(periode if periode in st.session_state.db_checklist.get(merek, {}).get(tipe, {}) else "Bulanan", {})
    hasil_checklist = []

    if menu == "Maintenance":
        if "Mati" in data_alat:
            st.markdown('<div class="section-title"> KONDISI ALAT MATI (Power OFF)</div>', unsafe_allow_html=True)
            with st.container(border=True):
                for i, step in enumerate(data_alat["Mati"]):
                    col1, col2 = st.columns([3, 2])
                    col1.write(f"**{i+1}.** {step}")
                    res = col2.radio("Status", ["Cleaned", "Not Cleaned"], key=f"m_{i}", horizontal=True, label_visibility="collapsed")
                    hasil_checklist.append({"Langkah": f"(Mati) {step}", "Hasil": res})
                    if i < len(data_alat["Mati"]) - 1: st.divider()
                    
        if "Hidup" in data_alat:
            st.markdown('<div class="section-title"> KONDISI ALAT HIDUP (Power ON)</div>', unsafe_allow_html=True)
            with st.container(border=True):
                for i, step in enumerate(data_alat["Hidup"]):
                    col1, col2 = st.columns([3, 2])
                    col1.write(f"**{i+1}.** {step}")
                    res = col2.radio("Status", ["Pass", "Fail", "⚠️ Warning"], key=f"h_{i}", horizontal=True, label_visibility="collapsed")
                    hasil_checklist.append({"Langkah": f"(Hidup) {step}", "Hasil": res})
                    if i < len(data_alat["Hidup"]) - 1: st.divider()
    else:
        st.markdown('<div class="section-title">🛠️ DETAIL TINDAKAN SERVICE</div>', unsafe_allow_html=True)
        data_service = st.session_state.db_checklist.get(merek, {}).get(tipe, {}).get("Service")
        if not data_service:
            template = st.session_state.db_checklist.get(merek, {}).get(tipe, {}).get("Bulanan", {})
            data_service = template.get("Mati", []) + template.get("Hidup", []) if isinstance(template, dict) else template

        with st.container(border=True):
            if data_service:
                for i, step in enumerate(data_service):
                    col1, col2 = st.columns([3, 2])
                    col1.write(f"**{i+1}.** {step}")
                    res = col2.radio("Hasil", ["Good", "⚠️Warning", "Replace"], key=f"svc_{i}", horizontal=True, label_visibility="collapsed")
                    hasil_checklist.append({"Langkah": f"(Service) {step}", "Hasil": res})
                    if i < len(data_service) - 1: st.divider()

    st.markdown('<div class="section-title">DOKUMENTASI & VALIDASI</div>', unsafe_allow_html=True)
    with st.container(border=True):
        catatan = st.text_area("Catatan Tambahan")
        foto_alat = st.file_uploader("Upload Foto Lapangan", accept_multiple_files=True, type=['jpg','png'])
        v1, v2 = st.columns(2)
        with v1: ttd_tek = st.file_uploader("TTD Teknisi", type=['png','jpg'], key="tek_ttd")
        with v2: ttd_cus = st.file_uploader("TTD Customer", type=['png','jpg'], key="cus_ttd")

    st.markdown("---")
    
    if 'proses_selesai' not in st.session_state:
        st.session_state.proses_selesai = False
        st.session_state.data_excel = None
        st.session_state.data_word = None
        st.session_state.nama_file = ""
    
    if st.button("➣ SUBMIT & GENERATE REPORT"):
        if not sn_alat or not nama_customer:
            st.error("⚠️ S/N Alat dan Nama Customer wajib diisi!")
        else:
            checklist_text = "\n".join([f"{c['Langkah']}: {c['Hasil']}" for c in hasil_checklist])
            data_laporan = {
                "Timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "Teknisi": nama_teknisi,
                "Customer": nama_customer,
                "Alat": f"{merek} {tipe}",
                "SN": sn_alat,
                "Periode": periode,
                "Catatan": catatan,
                "Detail_Checklist": checklist_text
            }
            
            with st.status("Memproses laporan...", expanded=True) as status:
                st.write("Mengirim data ke Google Sheets...")
                sheets_ok = save_to_sheets(data_laporan)
                
                if sheets_ok: st.write("✅ Data berhasil masuk ke Google Sheets!")
                else: st.write("❌ Gagal mengirim ke Sheets, tapi file lokal tetap dibuat...")
                
                if menu == "Maintenance":
                    if periode == "Bulanan": p_title = "MONTHLY MAINTENANCE"
                    elif periode == "3 Bulan": p_title = "QUARTERLY MAINTENANCE"
                    elif periode == "6 Bulan": p_title = "SEMI-ANNUAL MAINTENANCE"
                    elif periode == "12 Bulan": p_title = "ANNUAL MAINTENANCE"
                    else: p_title = "PERIODIC MAINTENANCE"
                else:
                    p_title = "SERVICE & REPAIR"

                clean_cus = nama_customer.replace(" ", "_") if nama_customer else "Customer"
                tgl_file = datetime.now().strftime('%d%m%y')
                
                nama_file_excel = f"{p_title.replace(' ', '_')}-{clean_cus}-{tgl_file}.xlsx"
                
                output_excel = BytesIO()
                with pd.ExcelWriter(output_excel, engine='xlsxwriter') as writer:
                    workbook  = writer.book
                    worksheet = workbook.add_worksheet('Report')
                    
                    navy_blue = '#1c3d73'
                    border_col = '#D1D5DB'
                    title_fmt = workbook.add_format({'bold': True, 'font_size': 18, 'font_color': navy_blue, 'align': 'center', 'valign': 'vcenter'})
                    type_fmt = workbook.add_format({'bold': True, 'font_size': 12, 'align': 'center', 'valign': 'vcenter', 'bg_color': '#E2E8F0', 'font_color': '#1E293B', 'border': 1, 'border_color': border_col})
                    label_fmt = workbook.add_format({'bold': True, 'bg_color': '#F3F4F6', 'border': 1, 'border_color': border_col, 'font_size': 10})
                    data_fmt = workbook.add_format({'border': 1, 'border_color': border_col, 'text_wrap': True, 'font_size': 10, 'valign': 'top'})
                    header_table = workbook.add_format({'bold': True, 'fg_color': navy_blue, 'font_color': 'white', 'border': 1, 'align': 'center', 'valign': 'vcenter'})
                    name_label_fmt = workbook.add_format({'bold': True, 'align': 'center', 'valign': 'top', 'top': 1})
                    
                    worksheet.set_column('A:A', 6)
                    worksheet.set_column('B:B', 50)
                    worksheet.set_column('C:C', 18)
                    worksheet.set_column('E:E', 20)
                    worksheet.set_column('F:F', 35)
                    
                    worksheet.merge_range('A1:F2', 'PT. RAJAERBA INDOCHEM', title_fmt)
                    worksheet.merge_range('A3:F3', f"{p_title} REPORT", type_fmt)
                    
                    info_pairs = [
                        ("No. Laporan", f"REI/{menu[:3].upper()}/{datetime.now().strftime('%Y%m%d')}/{sn_alat[:5].upper() if sn_alat else '000'}"),
                        ("Tanggal", datetime.now().strftime('%d %B %Y')),
                        ("Teknisi", nama_teknisi),
                        ("Customer", nama_customer),
                        ("Model Alat", f"{merek} {tipe}"),
                        ("S/N Number", sn_alat)
                    ]
                     
                    curr_row = 4
                    for label, val in info_pairs:
                        worksheet.write(curr_row, 4, label, label_fmt)
                        worksheet.write(curr_row, 5, val, data_fmt)
                        curr_row += 1
                    
                    worksheet.write(4, 0, 'NO', header_table)
                    worksheet.write(4, 1, 'DESKRIPSI PEKERJAAN', header_table)
                    worksheet.write(4, 2, 'STATUS', header_table)
                    
                    row_idx = 5
                    for i, item in enumerate(hasil_checklist):
                        bg = '#FFFFFF' if i % 2 == 0 else '#F9FAFB'
                        row_f = workbook.add_format({'border': 1, 'border_color': border_col, 'bg_color': bg, 'text_wrap': True})
                        worksheet.write(row_idx, 0, i+1, row_f)
                        worksheet.write(row_idx, 1, item['Langkah'], row_f)
                        worksheet.write(row_idx, 2, item['Hasil'], row_f)
                        row_idx += 1

                    sig_row = row_idx + 2
                    worksheet.set_row(sig_row + 1, 60)
                    
                    worksheet.merge_range(sig_row, 0, sig_row, 1, "Customer Approval,", workbook.add_format({'italic': True, 'align': 'center'}))
                    worksheet.write(sig_row, 2, "Technician,", workbook.add_format({'italic': True, 'align': 'center'}))
                    
                    sig_text_fmt = workbook.add_format({'italic': True, 'align': 'center', 'font_size': 9})
                    if ttd_cus: worksheet.insert_image(sig_row + 1, 0, "ttd_cus.png", {'image_data': ttd_cus, 'x_scale': 0.4, 'y_scale': 0.4, 'x_offset': 40})
                    else: worksheet.merge_range(sig_row + 1, 0, sig_row + 1, 1, "(Belum Ttd)", sig_text_fmt)
                    
                    if ttd_tek: worksheet.insert_image(sig_row + 1, 2, "ttd_tek.png", {'image_data': ttd_tek, 'x_scale': 0.4, 'y_scale': 0.4, 'x_offset': 10})
                    else: worksheet.write(sig_row + 1, 2, "(Belum Ttd)", sig_text_fmt)

                    worksheet.merge_range(sig_row + 2, 0, sig_row + 2, 1, f"( {nama_customer} )", name_label_fmt)
                    worksheet.write(sig_row + 2, 2, f"( {nama_teknisi} )", name_label_fmt)

                list_foto_input = [f for f in foto_alat] if foto_alat else []
                word_data = generate_docx(hasil_checklist, nama_teknisi, id_teknisi, nama_customer, sn_alat, catatan, merek, tipe, periode, nama_pt, lokasi_pt, list_foto_input, ttd_tek, ttd_cus)
                
                st.session_state.data_excel = output_excel.getvalue()
                st.session_state.data_word = word_data
                st.session_state.nama_file = f"{p_title.replace(' ', '_')}-{clean_cus}"
                st.session_state.proses_selesai = True
                
                status.update(label="Laporan Siap Diunduh!", state="complete", expanded=False)
                st.balloons()
        
    if st.session_state.proses_selesai:
        st.markdown('<div class="section-title">📥 DOWNLOAD AREA</div>', unsafe_allow_html=True)
        
        c_dl1, c_dl2 = st.columns(2)
        with c_dl1: st.download_button(label="📥 Download Excel", data=st.session_state.data_excel, file_name=f"{st.session_state.nama_file}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
        with c_dl2: st.download_button(label="📄 Download Word", data=st.session_state.data_word, file_name=f"{st.session_state.nama_file}.docx", use_container_width=True)
            
        no_wa_cs = "6281545939640"
        pesan_wa = f"*LAPORAN REI*\nTeknisi: {nama_teknisi}\nCustomer: {nama_customer}\nSN: {sn_alat}\nStatus: Laporan Laporan Selesai Terbuat."
        encoded_msg = urllib.parse.quote(pesan_wa)
        st.markdown(f'<a href="https://wa.me/{no_wa_cs}?text={encoded_msg}" target="_blank" class="whatsapp-btn">📲 Share ke WhatsApp CS</a>', unsafe_allow_html=True)
        
        if st.button("🔄 Buat Laporan Baru"):
            st.session_state.proses_selesai = False
            st.rerun()

# --- MODUL PENGATURAN USER-FRIENDLY (TANPA JSON RIBET) ---
elif menu == "Pengaturan":
    st.markdown('<div class="section-title">⚙️ Pengaturan Aplikasi</div>', unsafe_allow_html=True)
    
    # 1. WADAH ATAS: EDIT DATA CHECKLIST RAMAH PENGGUNA
    with st.container(border=True):
        st.subheader("📝 Edit Data Checklist Alat")
        st.write("Silakan Pilih Tipe Alat dan Tipe Form Pekerjaan yang Ingin di Edit Data Checklistnya:")
        
        # Pilihan berjenjang agar user mudah menembus data
        p_merek = st.selectbox("1. Pilih Merek Alat:", list(st.session_state.db_alat.keys()), key="p_mrk")
        p_tipe = st.selectbox("2. Pilih Tipe Alat:", st.session_state.db_alat[p_merek], key="p_tp")
        p_kondisi = st.selectbox("3. Pilih Jenis Tindakan Pekerjaan:", ["Mati (Power OFF)", "Hidup (Power ON)", "Service & Repair"], key="p_knd")
        
        # Mapping teks pilihan ke struktur database asli
        kondisi_map = {
            "Mati (Power OFF)": ("Bulanan", "Mati"),
            "Hidup (Power ON)": ("Bulanan", "Hidup"),
            "Service & Repair": ("Service", None)
        }
        
        kat_utama, kat_sub = kondisi_map[p_kondisi]
        
        # Mengambil data teks checklist yang ada saat ini
        if kat_sub:
            current_list = st.session_state.db_checklist.get(p_merek, {}).get(p_tipe, {}).get(kat_utama, {}).get(kat_sub, [])
        else:
            current_list = st.session_state.db_checklist.get(p_merek, {}).get(p_tipe, {}).get(kat_utama, [])
            
        # Jika data ternyata kosong/belum didefinisikan, buat list kosong
        if not isinstance(current_list, list):
            current_list = []
            
        # Gabungkan list menjadi teks bertingkat baris baru (seperti Notepad biasa)
        current_text_value = "\n".join(current_list)
        
        st.write("---")
        st.caption(f"✍️ **Panduan:** Edit teks di bawah ini. **Satu baris kalimat sama dengan satu poin checklist baru** di lembar form.")
        
        edited_text_lines = st.text_area(
            f"Daftar Langkah Kerja ({p_merek} - {p_tipe} [{kat_sub if kat_sub else kat_utama}]):",
            value=current_text_value,
            height=220,
            key="area_txt"
        )
        
        if st.button("💾 Simpan Perubahan Langkah Kerja"):
            # Pecah kembali teks baris baru menjadi list array murni Python
            new_list_data = [line.strip() for line in edited_text_lines.split("\n") if line.strip()]
            
            # Kembalikan struktur datanya ke session state secara presisi tanpa merusak struktur lain
            if p_merek not in st.session_state.db_checklist:
                st.session_state.db_checklist[p_merek] = {}
            if p_tipe not in st.session_state.db_checklist[p_merek]:
                st.session_state.db_checklist[p_merek][p_tipe] = {}
                
            if kat_sub:
                if kat_utama not in st.session_state.db_checklist[p_merek][p_tipe]:
                    st.session_state.db_checklist[p_merek][p_tipe][kat_utama] = {}
                st.session_state.db_checklist[p_merek][p_tipe][kat_utama][kat_sub] = new_list_data
            else:
                st.session_state.db_checklist[p_merek][p_tipe][kat_utama] = new_list_data
                
            st.success(f"🎉 Sukses! Checklist pekerjaan untuk {p_merek} {p_tipe} berhasil diperbarui.")

    st.write("") # Memberi jarak pemisah visual antar kontainer halaman

    # 2. WADAH BAWAH: DATA LINK-LINK GOOGLE DRIVE DLL
    with st.container(border=True):
        st.subheader("🔗 Data Dokumen & Link Google Drive Cloud")
        st.write("Atur daftar link referensi penting Google Drive atau URL eksternal pendukung di bawah ini:")
        
        updated_links = {}
        for key, value in st.session_state.link_drive_data.items():
            updated_links[key] = st.text_input(f"URL Tautan untuk {key}:", value=value)
            
        if st.button("💾 Amankan & Simpan Link Drive"):
            st.session_state.link_drive_data = updated_links
            st.success("✅ Semua tautan link awan Google Drive berhasil disimpan ke sistem!")
